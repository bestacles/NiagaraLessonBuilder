from docx import Document
from docx.shared import Pt, RGBColor, Inches

class LessonPlanDocGenerator:
    """
    Generates a Word (.docx) lesson plan document with half-inch margins,
    bold headings, bulleted lists with nested indentation based on leading spaces,
    and consistent spacing, including an Appendix. Omits empty sections.
    """
    def __init__(self, data):
        self.data = data
        self.doc = Document()
        # Half-inch margins
        sec = self.doc.sections[0]
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.5)
        sec.right_margin = Inches(0.5)
        self._set_styles()

    def _set_styles(self):
        normal = self.doc.styles['Normal']
        normal.font.name = 'Arial'
        normal.font.size = Pt(12)
        normal.font.color.rgb = RGBColor(0, 0, 0)
        pf = normal.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)

        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(6)

    def generate_docx(self, path):
        self._add_header()
        self._add_section_i()
        self._add_section_ii()
        self._add_section_iii()
        self._add_section_iv()
        self._add_section_v()
        self._add_section_vi()
        self._add_section_appendix()
        self.doc.save(path)

    def _add_header(self):
        p = self.doc.add_paragraph()
        r = p.add_run('NAME: '); r.bold = True
        p.add_run(self.data.get('name',''))
        r = p.add_run('    DATE: '); r.bold = True
        p.add_run(self.data.get('date',''))

        p = self.doc.add_paragraph()
        r = p.add_run('SUBJECT OF LESSON: '); r.bold = True
        p.add_run(self.data.get('subject',''))
        r = p.add_run('    GRADE LEVEL: '); r.bold = True
        p.add_run(self.data.get('grade',''))

        p = self.doc.add_paragraph()
        r = p.add_run('TIME ESTIMATE: '); r.bold = True
        p.add_run(self.data.get('time_estimate',''))
        r = p.add_run('    TOPIC: '); r.bold = True
        p.add_run(self.data.get('topic',''))

        self.doc.add_paragraph()

    def _add_section_i(self):
        slos = self.data.get('learning_objectives','').splitlines()
        targets = self.data.get('i_can_statements','').splitlines()
        exps = self.data.get('curriculum_expectations','').splitlines()
        # strip and filter
        slos = [l for l in slos if l.strip()]
        targets = [l for l in targets if l.strip()]
        exps = [l for l in exps if l.strip()]
        if not (slos or targets or exps): return
        self.doc.add_heading('I. STUDENT LEARNING OUTCOMES (SLOs)', level=1)
        for raw in slos:
            level = (len(raw) - len(raw.lstrip(' '))) // 2
            text = raw.strip()
            p = self.doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5 + 0.25 * level)
            p.add_run(text)
        if targets:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run('Learning Targets:'); run.bold = True
            for raw in targets:
                level = (len(raw) - len(raw.lstrip(' '))) // 2
                text = raw.strip()
                b = self.doc.add_paragraph(style='List Bullet')
                b.paragraph_format.left_indent = Inches(1 + 0.25 * level)
                b.add_run(text)
        if exps:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run('Ontario Curriculum Expectations:'); run.bold = True
            for raw in exps:
                level = (len(raw) - len(raw.lstrip(' '))) // 2
                text = raw.strip()
                b = self.doc.add_paragraph(style='List Bullet')
                b.paragraph_format.left_indent = Inches(1 + 0.25 * level)
                b.add_run(text)

    def _add_section_ii(self):
        sm = [l for l in self.data.get('student_materials','').splitlines() if l.strip()]
        tm = [l for l in self.data.get('teacher_materials','').splitlines() if l.strip()]
        pk = [l for l in self.data.get('prior_knowledge','').splitlines() if l.strip()]
        pa = [l for l in self.data.get('prep_ahead','').splitlines() if l.strip()]
        if not (sm or tm or pk or pa): return
        self.doc.add_heading('II. PREPARATION', level=1)
        if sm or tm:
            p = self.doc.add_paragraph('Materials:')
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].bold = True
            for raw in sm + tm:
                level = (len(raw) - len(raw.lstrip(' '))) // 2
                text = raw.strip()
                b = self.doc.add_paragraph(style='List Bullet')
                b.paragraph_format.left_indent = Inches(1 + 0.25 * level)
                b.add_run(text)
        if pk:
            p = self.doc.add_paragraph('Prior Knowledge:')
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].bold = True
            for raw in pk:
                level = (len(raw) - len(raw.lstrip(' '))) // 2
                text = raw.strip()
                b = self.doc.add_paragraph(style='List Bullet')
                b.paragraph_format.left_indent = Inches(1 + 0.25 * level)
                b.add_run(text)
        if pa:
            p = self.doc.add_paragraph('Need to Do Ahead of Time:')
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].bold = True
            for raw in pa:
                level = (len(raw) - len(raw.lstrip(' '))) // 2
                text = raw.strip()
                b = self.doc.add_paragraph(style='List Bullet')
                b.paragraph_format.left_indent = Inches(1 + 0.25 * level)
                b.add_run(text)

    def _add_section_iii(self):
        sections = [
            ('A. Anticipatory Set / Opening:', 'anticipatory'),
            ('B1. Activity #1 (Mini Lesson):', 'activity_1'),
            ('B2. Activity #2 (Guided Practice):', 'activity_2'),
            ('B3. Activity #3 (Independent/Formal Assessment):', 'activity_3'),
            ('C. Technology Integration:', 'technology'),
            ('D. Unfinished Work / Homework:', 'unfinished'),
            ('E. Alternative Plan:', 'alt_plan'),
        ]
        any_content = False
        for title, key in sections:
            raw_items = self.data.get(key,'' ).splitlines()
            items = [l for l in raw_items if l.strip()]
            if not items: continue
            if not any_content:
                self.doc.add_heading('III. BODY OF THE LESSON', level=1)
                any_content = True
            p = self.doc.add_paragraph(title)
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].bold = True
            for raw in items:
                level = (len(raw) - len(raw.lstrip(' '))) // 2
                text = raw.strip()
                b = self.doc.add_paragraph(style='List Bullet')
                b.paragraph_format.left_indent = Inches(1 + 0.25 * level)
                b.add_run(text)

    def _add_section_iv(self):
        groups = [
            ('Struggling Students', self.data.get('struggle_1',''), self.data.get('struggle_2','')),
            ('Language Learners', self.data.get('ell_1',''), self.data.get('ell_2','')),
            ('Gifted & Talented Students', self.data.get('gifted_1',''), self.data.get('gifted_2','')),
        ]
        flat = []
        for label, a, b in groups:
            flat.extend([l for l in (a, b) if l.strip()])
        if not flat: return
        self.doc.add_heading('IV. ACCOMMODATIONS', level=1)
        for label, a, b in groups:
            items = [l for l in (a, b) if l.strip()]
            if not items: continue
            p = self.doc.add_paragraph(label + ':')
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].bold = True
            for raw in items:
                level = (len(raw) - len(raw.lstrip(' '))) // 2
                text = raw.strip()
                b = self.doc.add_paragraph(style='List Bullet')
                b.paragraph_format.left_indent = Inches(1 + 0.25 * level)
                b.add_run(text)

    def _add_section_v(self):
        lines = [l for l in self.data.get('closure','').splitlines() if l.strip()]
        if not lines: return
        self.doc.add_heading('V. CLOSURE', level=1)
        for raw in lines:
            level = (len(raw) - len(raw.lstrip(' '))) // 2
            text = raw.strip()
            b = self.doc.add_paragraph(style='List Bullet')
            b.paragraph_format.left_indent = Inches(0.5 + 0.25 * level)
            b.add_run(text)

    def _add_section_vi(self):
        fa = [l for l in self.data.get('formal_assessment','').splitlines() if l.strip()]
        ia = [l for l in self.data.get('informal_assessment','').splitlines() if l.strip()]
        if not (fa or ia): return
        self.doc.add_heading('VI. ASSESSMENT', level=1)
        if fa:
            p = self.doc.add_paragraph()
            run = p.add_run('Formal Assessment:'); run.bold = True
            p.paragraph_format.left_indent = Inches(0.5)
            for raw in fa:
                level = (len(raw) - len(raw.lstrip(' '))) // 2
                text = raw.strip()
                b = self.doc.add_paragraph(style='List Bullet')
                b.paragraph_format.left_indent = Inches(1 + 0.25 * level)
                b.add_run(text)
        if ia:
            p2 = self.doc.add_paragraph()
            run2 = p2.add_run('Informal Assessment:'); run2.bold = True
            p2.paragraph_format.left_indent = Inches(0.5)
            for raw in ia:
                level = (len(raw) - len(raw.lstrip(' '))) // 2
                text = raw.strip()
                b2 = self.doc.add_paragraph(style='List Bullet')
                b2.paragraph_format.left_indent = Inches(1 + 0.25 * level)
                b2.add_run(text)

    def _add_section_appendix(self):
        files = self.data.get('appendix_files', [])
        refs  = [l for l in self.data.get('apa_references','').splitlines() if l.strip()]
        if not (files or refs): return
        self.doc.add_heading('VII. APPENDIX', level=1)
        for f in files:
            b = self.doc.add_paragraph(style='List Bullet')
            b.paragraph_format.left_indent = Inches(1)
            b.add_run(f)
        if refs:
            p = self.doc.add_paragraph()
            run = p.add_run('References:'); run.bold = True
            p.paragraph_format.left_indent = Inches(0.5)
            for raw in refs:
                level = (len(raw) - len(raw.lstrip(' '))) // 2
                text = raw.strip()
                b = self.doc.add_paragraph(style='List Bullet')
                b.paragraph_format.left_indent = Inches(1 + 0.25 * level)
                b.add_run(text)
